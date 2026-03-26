import streamlit as st
from docx import Document
from docx.shared import Inches
import tempfile
import re

st.set_page_config(page_title="SISGÉN CORE", layout="centered")

# -------- LOGIN --------
if "auth" not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    st.title("Acceso SISGÉN")

    u = st.text_input("Usuario")
    p = st.text_input("Contraseña", type="password")

    if st.button("Ingresar"):
        if u == "admin" and p == "1234":
            st.session_state.auth = True
        else:
            st.error("Credenciales incorrectas")

    st.stop()

# -------- UI --------
st.title("Motor Documental SISGÉN")

empresa = st.text_input("Empresa")
representante = st.text_input("Representante")
correo = st.text_input("Correo")
fecha = st.text_input("Fecha")

logo = st.file_uploader("Logo", type=["png","jpg"])
archivo = st.file_uploader("Documento", type=["docx"])

# -------- FUNCIÓN SEGURA --------
def reemplazar_run(run):
    texto = run.text

    if not texto:
        return

    if "@" in texto:
        run.text = correo

    elif re.search(r"\d{1,2}/\d{1,2}/\d{2,4}", texto):
        run.text = fecha

    elif texto.isupper() and len(texto.split()) > 1:
        run.text = representante

    elif any(x in texto.lower() for x in ["conjunto","ph","residencial"]):
        run.text = empresa


# -------- PROCESO --------
if st.button("Generar"):

    try:

        if not archivo:
            st.warning("Sube documento")
            st.stop()

        doc = Document(archivo)

        # TEXTO
        for p in doc.paragraphs:
            for run in p.runs:
                reemplazar_run(run)

        # TABLAS
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for run in p.runs:
                            reemplazar_run(run)

        # LOGO SEGURO
        if logo:
            try:
                temp_logo = tempfile.NamedTemporaryFile(delete=False)
                temp_logo.write(logo.read())

                for s in doc.sections:
                    header = s.header
                    if header.paragraphs:
                        run = header.paragraphs[0].add_run()
                        run.add_picture(temp_logo.name, width=Inches(1.2))

            except:
                st.warning("Logo no compatible")

        # PIE
        for s in doc.sections:
            footer = s.footer
            footer.paragraphs[0].text = f"Fecha generación: {fecha}"

        # GUARDAR
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)

        st.success("Documento listo")

        with open(tmp.name, "rb") as f:
            st.download_button("Descargar", f)

    except Exception as e:
        st.error("Error controlado")
        st.write(e)
